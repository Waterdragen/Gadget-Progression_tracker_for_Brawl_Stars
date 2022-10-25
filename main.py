from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt, RGBColor, Mm


def write_docx():
    global BoxStats, ProgressionStats, BrawlPassStats, ClubAssets, StarPointsBox, Personal, MaxAccReq
    global BpGains, ResPerSsn, CcEqResPerSsn, PP2Coins, TotalTime, SpAccounting

    document = Document()
    section = document.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.left_margin = Mm(25.4)
    section.right_margin = Mm(25.4)
    section.top_margin = Mm(25.4)
    section.bottom_margin = Mm(25.4)
    paragraph = document.add_paragraph()
    blue = RGBColor(60, 120, 215)
    magenta = RGBColor(183, 38, 215)
    green = RGBColor(56, 118, 30)
    brown = RGBColor(166, 28, 1)

    def add_run(content, font_name="Oxygen Mono", f_size=11, f_bold=False, color=RGBColor(0, 0, 0)):
        nonlocal paragraph
        _run = paragraph.add_run(content)
        _run.font.name = font_name
        _run.font.size = Pt(f_size)
        _run.font.bold = f_bold
        _run.font.color.rgb = color

    # ================================================Box stats & assumptions
    add_run("Box assets & assumptions\n\n", f_size=14)
    add_run("Max brawlers: ")
    add_run(f"{BoxStats[0]}\n", color=blue, f_bold=True)
    add_run("100-box sampling size\n\tCoins: ")
    add_run(f"{BoxStats[1]:.3f} ", color=blue, f_bold=True)
    add_run("per box\n\tPP: ")
    add_run(f"{BoxStats[2]:.3f} ", color=blue, f_bold=True)
    add_run("per box\nGadget: ")
    add_run(f"1 per {1 / BoxStats[3]:.3f} ", color=blue, f_bold=True)
    add_run("boxes\nStar power: ")
    add_run(f"1 per {1 / BoxStats[4]:.3f} ", color=blue, f_bold=True)
    add_run("boxes\nToken doublers: ")
    add_run(f"{BoxStats[5]:.3f} ", color=blue, f_bold=True)
    add_run("per box, each 200 tokens\nHas unlocked all brawlers\nGadgets & Star powers coexist\n"
            "Club league rank: Average of Mythic I to III indiv. max rewards\nMax rewards in club games\n"
            "Free to play, all resources on progression\n5000 star points spent/ 33 boxes gained biweekly\n"
            "0 challenge rewards\nInstantly gets tokens once available (why calculate play time?)\n\n")

    # ================================================Pre-calculation
    add_run("Pre-calculation\n\n", f_size=14)
    add_run("Coins to max 1 brawler: sum([20, 35, 75, 140, 290, 480, 800, 1250, 1875, 2800]) = 7765\n"
            "PP to max 1 brawler: sum([20, 30, 50, 80, 130, 210, 340, 550, 890, 1440]) = 3740\n"
            "Gear costs in coins: {“super rare”: 1000, “epic”: 1500, “mythic”: 2000}\n\n"
            "\tMax account requirements")
    if Personal[0] is None:
        add_run(f"Coins to max PP: 7765 * {BoxStats[0]} = {MaxAccReq[0]}\n")
    else:
        add_run(f"Coins to max PP: {Personal[0]}\n")
    super_rare_gear = ProgressionStats[2] if Personal[4] is None else Personal[4]
    epic_gear = ProgressionStats[3] if Personal[5] is None else Personal[5]
    mythic_gear = ProgressionStats[4] if Personal[6] is None else Personal[6]
    add_run(f"Coins to max gears: ")
    if Personal[4] is None:
        add_run(f"{super_rare_gear} ")
    else:
        add_run(f"{super_rare_gear} ", color=magenta, f_bold=True)
    add_run("* 1000 + ")
    if Personal[5] is None:
        add_run(f"{epic_gear} ")
    else:
        add_run(f"{epic_gear} ", color=magenta, f_bold=True)
    add_run("* 1500 + ")
    if Personal[6] is None:
        add_run(f"{mythic_gear} ")
    else:
        add_run(f"{mythic_gear} ", color=magenta, f_bold=True)
    add_run(f"* 2000 = {MaxAccReq[1]}\n")
    if Personal[0] is None:
        add_run(f"Coins to max ALL: {MaxAccReq[0]} + {MaxAccReq[1]} = ")
        add_run(f"{MaxAccReq[2]}\n", color=blue, f_bold=True)
    else:
        add_run("Coins to max ALL: ")
        add_run(f"{Personal[0]}\n", color=magenta, f_bold=True)
    if Personal[1] is None:
        add_run(f"PP to max ALL: 3740 * {BoxStats[0]} = ")
        add_run(f"{MaxAccReq[3]}\n", color=blue, f_bold=True)
    else:
        add_run("PP to max ALL: ")
        add_run(f"{Personal[1]}\n", color=magenta, f_bold=True)
    if Personal[2] is None:
        add_run(f"Gadgets to max ALL: 2 * {BoxStats[0]} = {MaxAccReq[4]}\n")
    else:
        add_run("Gadgets to max ALL: ")
        add_run(f"{Personal[2]}\n", color=magenta, f_bold=True)
    if Personal[3] is None:
        add_run(f"Star powers to max ALL: {MaxAccReq[5]}\n\n")
    else:
        add_run("Star powers to max ALL: ")
        add_run(f"{Personal[3]}\n\n", color=magenta, f_bold=True)
    add_run(f"**Star powers that have not been obtained will be considered later.\n\n")

    # ================================================Brawl Pass Calculation
    add_run("Brawl Pass Calculation\n\n", f_size=14)
    add_run(f"Brawl pass duration: 52 / 6 * 7 = ")
    add_run("60.667d\n", color=blue, f_bold=True)
    add_run(f"Token generation: 200/d + 8 * 5 new event = 240/d\n"
            f"Free seasonal quest tokens: (250 * 3 + 500 * 3)/wk = 321.429/d\n"
            f"Premium seasonal quest tokens: (250 * 4 + 500 * 4)/wk = 428.571/d\n"
            f"Special event tokens: 500/wk = 71.429/d\n\n"
            f"Quest tokens Free (default + daily + season + special):\n\t"
            f"{BrawlPassStats[1]}/d + 2 * {BrawlPassStats[2]/2}/d + {BrawlPassStats[3]/7:.3f}/d "
            f"+ {BrawlPassStats[5]/7:.3f}/d = ")
    add_run(f"{BpGains[0]:.3f}\n", color=blue, f_bold=True)
    add_run(f"Quest tokens Premium:\n\t{BrawlPassStats[1]}/d + 2 * {BrawlPassStats[2]/2}/d + "
            f"{BrawlPassStats[4]/7:.3f}/d + {BrawlPassStats[5]/7:.3f}/d = ")
    add_run(f"{BpGains[1]:.3f}\n\n", color=blue, f_bold=True)

    add_run(f"Tokens to tier 70: {BrawlPassStats[6]}\n"
            f"Brawl Pass per season: 90 / 169 = ")
    add_run(f"{BrawlPassStats[7]:.3f}\n", color=blue, f_bold=True)
    add_run(f"Free track boxes: {BrawlPassStats[8]}\n")
    add_run(f"Full BP boxes: {BrawlPassStats[9]}\n")
    add_run(f"Free track coins: {BrawlPassStats[10]}\n")
    add_run(f"Full BP coins: {BrawlPassStats[11]}\n")
    add_run(f"Free track PP: {BrawlPassStats[12]}\n")
    add_run(f"Full BP PP: {BrawlPassStats[13]}\n")
    add_run("Free track max bonus boxes:\n\t(")
    add_run(f"{BrawlPassStats[0]:.3f}d ", color=blue, f_bold=True)
    add_run("* ")
    add_run(f"{BpGains[0]:.3f}/d ", color=blue, f_bold=True)
    add_run(f" - {BrawlPassStats[6]}) // 500 * 3 = {BpGains[2]}\n")
    add_run("Premium max bonus boxes:\n\t(")
    add_run(f"{BrawlPassStats[0]:.3f}d ", color=blue, f_bold=True)
    add_run("* ")
    add_run(f"{BpGains[1]:.3f}/d ", color=blue, f_bold=True)
    add_run(f" - {BrawlPassStats[6]}) // 500 * 3 = {BpGains[3]}\n\n")

    add_run(f"Free full pass boxes: {BrawlPassStats[8]} + {BpGains[2]} = {BpGains[4]}\n")
    add_run(f"Premium full pass boxes: {BrawlPassStats[9]} + {BpGains[3]} = {BpGains[5]}\n")
    add_run(f"Star points boxes per season: 33 / 14 * ")
    add_run(f"{BrawlPassStats[0]:.3f}d ", color=blue, f_bold=True)
    add_run(f"= {StarPointsBox}\n")
    add_run("Bonus boxes per box by doublers: ")
    add_run(f"{BoxStats[5]:.3f} * 200 / 500 * 3 = {BpGains[6]:.3f}\n\n")

    anti_pass_ratio = 1 - BrawlPassStats[7]
    ave_coins = BrawlPassStats[10] * anti_pass_ratio + BrawlPassStats[11] * BrawlPassStats[7]
    ave_pp = BrawlPassStats[12] * anti_pass_ratio + BrawlPassStats[13] * BrawlPassStats[7]
    add_run(f"Average coins per season: {BrawlPassStats[10]} * {anti_pass_ratio:.3f} + "
            f"{BrawlPassStats[11]} * ")
    add_run(f"{BrawlPassStats[7]:.3f} ", color=blue, f_bold=True)
    add_run(f"= {ave_coins:.3f}\n")
    add_run(f"Average PP per season: {BrawlPassStats[12]} * {anti_pass_ratio:.3f} + {BrawlPassStats[13]} * ")
    add_run(f"{BrawlPassStats[7]:.3f} ", color=blue, f_bold=True)
    add_run(f"= {ave_pp:.3f}\n")
    add_run(f"Average boxes per season: {BpGains[4]} * {anti_pass_ratio:.3f} + {BpGains[5]} * ")
    add_run(f"{BrawlPassStats[7]:.3f} ", color=blue, f_bold=True)
    add_run(f"= {ResPerSsn[0]:.3f}\n")
    add_run(f"Average boxes per season with token doublers"
            f"(sum to inf. since doublers gained are converted to boxes):\n\t"
            f"{ResPerSsn[0]:.3f} / (1 - {BpGains[6]:.3f}) = ")
    add_run(f"{ResPerSsn[1]:.3f}\n", color=green, f_bold=True)
    add_run(f"Total equivalent coins per season:\n\t{ave_coins:.3f} + ")
    add_run(f"{ResPerSsn[1]:.3f} ", color=green)
    add_run("* ")
    add_run(f"{BoxStats[1]} ", color=blue)
    add_run("= ")
    add_run(f"{ResPerSsn[2]:.3f}\n", color=green, f_bold=True)
    add_run(f"Total equivalent PP per season:\n\t{ave_pp:.3f} + ")
    add_run(f"{ResPerSsn[1]:.3f} ", color=green)
    add_run("* ")
    add_run(f"{BoxStats[2]} ", color=blue)
    add_run("= ")
    add_run(f"{ResPerSsn[3]:.3f}\n\n", color=green, f_bold=True)

    # ================================================Club coins
    add_run("Club coins\n", f_size=14)
    if ClubAssets[0] == 682.792 and ClubAssets[1] == 652.667:
        add_run("\tAssuming Mythic rank\n", f_size=12)
        add_run(f"Average club league reward: mean([*]) = 682.792\n"
                f"Average club quest reward: mean([623, 653, 682]) = 652.667\n"
                f"Average club coins per week: mean([682.792, 652.667]) = 667.729\n"
                f"Average club coins per season: 667.729 / 7 * ")
    else:
        ave_cc = (ClubAssets[0] + ClubAssets[1]) / 2
        add_run(f"Average club league reward: {ClubAssets[0]}\n"
                f"Average club quest reward: {ClubAssets[1]}\n"
                f"Average club coins per week: mean([{ClubAssets[0]}, {ClubAssets[1]}]) = {ave_cc:.3f}\n"
                f"Average club coins per season: {ave_cc:.3f} / 7 * ")
    add_run(f"{BrawlPassStats[0]:.3f} ", color=blue, f_bold=True)
    add_run("= ")
    add_run(f"{CcEqResPerSsn[0]:.3f}\n", color=brown, f_bold=True)
    add_run("Equivalent coins: ")
    add_run(f"{CcEqResPerSsn[0]:.3f} ", color=brown)
    add_run("/ 2 * 5 = ")
    add_run(f"{CcEqResPerSsn[1]:.3f}\n", color=brown, f_bold=True)
    add_run("Equivalent PP: ")
    add_run(f"{CcEqResPerSsn[0]:.3f} ", color=brown)
    add_run("/ 3 * 5 = ")
    add_run(f"{CcEqResPerSsn[2]:.3f}\n\n", color=brown, f_bold=True)

    # ================================================Convert PP to coins
    add_run("Convert PP to coins\n", f_size=14)
    add_run("PP from boxes per season: ")
    add_run(f"{ResPerSsn[1]:.3f} ", color=green)
    add_run("* ")
    add_run(f"{BoxStats[2]} ", color=blue)
    add_run(f"= {PP2Coins[0]:.3f}\n")
    add_run(f"PP from pass per season: {PP2Coins[1]:.3f}\n")
    add_run("PP to coins ratio: ")
    add_run(f"{BoxStats[2]} ", color=blue)
    add_run("* 2 / ")
    add_run(f"{BoxStats[1]} ", color=blue)
    add_run(f"* {PP2Coins[1]:.3f}/({PP2Coins[0]:.3f}+{PP2Coins[1]:.3f}) + 2 * "
            f"{PP2Coins[1]:.3f}/({PP2Coins[0]:.3f}+{PP2Coins[1]:.3f}) = ")
    add_run(f"{PP2Coins[2]:.3f}\n\n", color=blue, f_bold=True)

    # ================================================Total Time Calculation
    add_run("Total Time Calculation\n\n", f_size=14)
    add_run("Seasons to max out coins: ")
    add_run(f"{MaxAccReq[2]} ", color=blue)
    add_run("/ ")
    add_run(f"{ResPerSsn[2]:.3f} ", color=green)
    add_run(f"= {TotalTime[0]:.3f}\n")
    add_run("Seasons to max out PP: ")
    add_run(f"{MaxAccReq[3]}", color=blue)
    add_run("/ (")
    add_run(f"{ResPerSsn[3]:.3f} ", color=green)
    add_run("+ ")
    add_run(f"{CcEqResPerSsn[2]:.3f}", color=brown)
    add_run(f") = {TotalTime[1]:.3f}\n")
    season_delta = TotalTime[0] - TotalTime[1]
    add_run(f"Seasons for coins exceed PP: {TotalTime[0]:.3f} - {TotalTime[1]:.3f} "
            f"= {season_delta:.3f}\n")
    add_run(f"Total time required (ignoring SP): {TotalTime[1]:.3f} + {season_delta:.3f} / (")
    add_run(f"{PP2Coins[2]:.3f} ", color=blue)
    add_run("+ 1) = ")
    add_run(f"{TotalTime[2]:.3f} ", color=blue, f_bold=True)
    add_run(f"seasons = {TotalTime[3]}\n\n")

    # ================================================Taking Star powers into account
    add_run("\tTaking Star powers into account\n", f_size=12)
    add_run("No. of boxes ever opened: ")
    boxes_lifetime = ResPerSsn[1] * TotalTime[2]
    add_run(f"{ResPerSsn[1]:.3f} ", color=green)
    add_run("* ")
    add_run(f"{TotalTime[2]:.3f} ", color=blue)
    add_run(f"= {boxes_lifetime:.3f}\n")
    add_run(f"No. of Star powers gained: floor({boxes_lifetime:.3f} * ")
    sp_owned = MaxAccReq[5] - SpAccounting[0]
    add_run(f"{BoxStats[4]}", color=blue)
    add_run(f") = {sp_owned}\n")
    add_run(f"Missing Star powers: {MaxAccReq[5]:.3f} - {sp_owned} = {SpAccounting[0]}\n")
    add_run("1 Box to Star power: (")
    add_run(f"{BoxStats[1]} ", color=blue)
    add_run("+ ")
    add_run(f"{BoxStats[2]} ", color=blue)
    add_run(") * 2 / 2000 + ")
    add_run(f"{BoxStats[4]} ", color=blue)
    add_run(f"= {SpAccounting[1]:.3f}\n")
    add_run(f"No. of coins required = {SpAccounting[0]} * 2000 = {SpAccounting[2]}\n")
    add_run(f"No. of coins required = {SpAccounting[0]} / {SpAccounting[1]:.3f} = {SpAccounting[3]:.3f}\n\n")

    add_run("Scenario 1 - pay 100% coins with club coins:\n\tExtra no. of seasons: ")
    scen_1 = SpAccounting[2] / CcEqResPerSsn[1]
    add_run(f"{SpAccounting[2]} / ")
    add_run(f"{CcEqResPerSsn[1]:.3f} ", color=brown)
    add_run(f"= {scen_1:.3f}\n")
    add_run("Scenario 2 - get SP and coins from boxes only:\n\tExtra no. of seasons: ")
    scen_2 = SpAccounting[3] / ResPerSsn[1]
    add_run(f"{SpAccounting[3]:.3f} / ")
    add_run(f"{ResPerSsn[1]:.3f} ", color=green)
    add_run(f"= {scen_2:.3f}\n")
    add_run(f"Scenario 1 + 2:\n\tExtra no. of seasons: ")
    add_run(f"1 / (1 / {scen_1:.3f} + 1 / {scen_2:.3f}) = ")
    add_run(f"{SpAccounting[4]:.3f}\n\n\n", color=blue, f_bold=True)

    add_run(f"Grand total time to max account:\n\t{TotalTime[2]:.3f} + {SpAccounting[4]:.3f} = ")
    grand_total_ss = TotalTime[2] + SpAccounting[4]
    grand_total_yr = grand_total_ss * BrawlPassStats[0] / 365
    add_run(f"{grand_total_ss:.3f} ", color=blue, f_bold=True)
    add_run("seasons = ")
    add_run(f"{grand_total_yr:.3f} ", color=blue, f_bold=True)
    add_run("years\n\n\n")

    # ================================================Colored Text
    add_run("Colored text\n", f_size=12)
    add_run("0 ", color=blue, f_bold=True)
    add_run("Important numbers\n", color=blue)
    add_run("1 ", color=magenta, f_bold=True)
    add_run("Personal overrides\n", color=magenta)
    add_run("2 ", color=green, f_bold=True)
    add_run("Resources (coins, PP, etc.) earned per season\n", color=green)
    add_run("3 ", color=brown, f_bold=True)
    add_run("Club coins equivalent resources (coins, PP, etc.) earned per season\n", color=brown)

    try:
        document.save("progression_report.docx")
    except PermissionError:
        raise SystemExit("Please close the report first.\nWriting aborted.")
    else:
        print("Saved as progression_report.docx")
    finally:
        print("Exiting program")


if __name__ == "__main__":
    Wb = load_workbook("Gadget - Progression tracker for Brawl Stars.xlsx", data_only=True)
    Ws = Wb["Time calculation"]
    BoxStats = [Ws[f'B{i + 3}'].value for i in range(6)]
    ProgressionStats = [Ws[f'B{i + 10}'].value for i in range(5)]
    BrawlPassStats = [Ws[f'B{i + 16}'].value for i in range(14)]
    ClubAssets = [Ws[f'B{i + 31}'].value for i in range(4)]
    StarPointsBox = Ws['B36'].value
    Personal = [Ws[f'B{i + 38}'].value for i in range(7)]
    MaxAccReq = [Ws[f'E{i + 3}'].value for i in range(6)]
    BpGains = [Ws[f'E{i + 10}'].value for i in range(7)]
    ResPerSsn = [Ws[f'E{i + 18}'].value for i in range(4)]
    CcEqResPerSsn = [Ws[f'E{i + 23}'].value for i in range(3)]
    PP2Coins = [Ws[f'E{i + 27}'].value for i in range(3)]
    TotalTime = [Ws[f'H{i + 3}'].value for i in range(4)]
    SpAccounting = [Ws[f'H{i + 8}'].value for i in range(5)]
    write_docx()
